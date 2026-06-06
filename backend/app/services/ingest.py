"""
DX-RAG 文件处理模块
支持多格式文件的文本提取、清洗和智能切片
"""
import base64
import io
import os
import re
from pathlib import Path
from typing import List, Optional

from app.core.config import settings


# ============================================================
#  第一部分：文件文本提取
# ============================================================

def extract_text_from_txt(file_path: Path) -> str:
    """从文本文件提取内容，自动处理编码"""
    encodings = ["utf-8", "utf-16", "gbk", "latin-1"]
    for enc in encodings:
        try:
            return file_path.read_text(encoding=enc)
        except (UnicodeDecodeError, UnicodeError):
            continue
    # 最后的兜底方案
    return file_path.read_text(encoding="gbk", errors="ignore")


def extract_text_from_pdf(file_path: Path) -> str:
    """从 PDF 文件提取文本，图片型 PDF 自动调用视觉模型"""
    text = ""

    # 方案一：使用 PyMuPDF（fitz）提取文本
    try:
        import fitz
        doc = fitz.open(str(file_path))
        parts = []
        for page in doc:
            page_text = page.get_text()
            if page_text:
                parts.append(page_text)
        doc.close()
        text = "\n".join(parts)
    except Exception:
        pass

    # 如果 PyMuPDF 未提取到文本，尝试 PyPDF2
    if not text.strip():
        try:
            from pypdf import PdfReader
            reader = PdfReader(str(file_path))
            parts = []
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    parts.append(page_text)
            text = "\n".join(parts)
        except Exception:
            pass

    # 图片型 PDF：文本为空时调用 DeepSeek Vision 视觉模型
    if not text.strip():
        text = _extract_text_with_deepseek_vision(str(file_path))

    return text


def _extract_text_with_deepseek_vision(pdf_path: str) -> str:
    """
    使用 DeepSeek Vision 模型提取图片型 PDF 中的文字
    DeepSeek 兼容 OpenAI 视觉 API，直接传 base64 图片
    """
    import fitz
    from openai import OpenAI

    api_key = settings.deepseek_api_key

    if not api_key or api_key.startswith("your-"):
        return ""  # 未配置 API Key，跳过

    client = OpenAI(
        api_key=api_key,
        base_url=settings.deepseek_base_url,
    )

    doc = fitz.open(pdf_path)
    full_text_parts = []

    for page_num in range(len(doc)):
        page = doc[page_num]

        # 将页面渲染为图片
        pix = page.get_pixmap(dpi=150)
        img_bytes = pix.tobytes("jpg")
        img_base64 = base64.b64encode(img_bytes).decode("utf-8")

        try:
            response = client.chat.completions.create(
                model="deepseek-chat",
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:image/jpeg;base64,{img_base64}"
                                },
                            },
                            {
                                "type": "text",
                                "text": "请提取图片中的所有文字，保持原始格式和排版。只输出提取到的文字内容，不要添加任何额外说明。",
                            },
                        ],
                    }
                ],
                max_tokens=4000,
            )

            content = response.choices[0].message.content
            if content:
                full_text_parts.append(content.strip())

        except Exception as e:
            print(f"[DeepSeek Vision] 第 {page_num + 1} 页提取失败: {e}")
            continue

    doc.close()
    return "\n\n".join(full_text_parts)


def extract_text_from_docx(file_path: Path) -> str:
    """从 Word 文档提取文本"""
    from docx import Document
    doc = Document(str(file_path))

    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    # 同时提取表格中的文本
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text for cell in row.cells if cell.text.strip())
            if row_text.strip():
                parts.append(row_text)

    return "\n".join(parts)


def extract_text_from_xlsx(file_path: Path) -> str:
    """从 Excel 表格提取文本"""
    from openpyxl import load_workbook

    wb = load_workbook(str(file_path), data_only=True)
    parts = []

    for sheet in wb.worksheets:
        sheet_name = sheet.title
        row_texts = []
        for row in sheet.iter_rows(values_only=True):
            row_text = " ".join(str(cell) for cell in row if cell is not None)
            if row_text.strip():
                row_texts.append(row_text)

        if row_texts:
            parts.append(f"--- 工作表: {sheet_name} ---")
            parts.extend(row_texts)

    return "\n".join(parts)


def read_text_from_file(file_path: Path) -> str:
    """
    统一的文件文本提取入口

    Args:
        file_path: 文件路径

    Returns:
        提取的文本内容
    """
    suffix = file_path.suffix.lower()

    # 文本类文件
    if suffix in {".txt", ".md", ".csv", ".json", ".log"}:
        return extract_text_from_txt(file_path)

    # PDF 文件
    if suffix == ".pdf":
        return extract_text_from_pdf(file_path)

    # Word 文件
    if suffix == ".docx":
        return extract_text_from_docx(file_path)

    # Excel 文件
    if suffix in {".xlsx", ".xlsm", ".xltx", ".xltm"}:
        return extract_text_from_xlsx(file_path)

    raise ValueError(f"不支持的文件格式: {suffix}")


# ============================================================
#  第二部分：文本清洗
# ============================================================

def clean_text(text: str) -> str:
    """
    文本清洗：去除空行、统一编码、基础噪声过滤

    Args:
        text: 原始文本

    Returns:
        清洗后的文本
    """
    if not text:
        return ""

    # 按行处理
    cleaned_lines = []
    for line in text.splitlines():
        stripped = line.strip()

        # 跳过空行
        if not stripped:
            continue

        # 跳过纯符号行（噪声）
        if re.match(r'^[\s\-_=#*~\.…]{5,}$', stripped):
            continue

        # 保留有意义的内容
        cleaned_lines.append(stripped)

    # 重新拼接
    cleaned = "\n".join(cleaned_lines)

    # 统一多余空白
    cleaned = re.sub(r'\n{3,}', '\n\n', cleaned)  # 最多保留一个空行
    cleaned = re.sub(r' {2,}', ' ', cleaned)       # 多余空格合并

    return cleaned.strip()


# ============================================================
#  第三部分：文本切片
# ============================================================

def split_text_by_headers(text: str) -> List[str]:
    """
    Markdown 标题切分：按标题层次结构切分文本

    Args:
        text: 输入文本

    Returns:
        切分后的文本块列表
    """
    from langchain_text_splitters import MarkdownHeaderTextSplitter

    headers_to_split = [
        ("#", "大章节"),
        ("##", "小节"),
        ("###", "小点"),
        ("####", "段落"),
    ]

    try:
        header_splitter = MarkdownHeaderTextSplitter(
            headers_to_split_on=headers_to_split,
            strip_headers=False,
        )
        docs = header_splitter.split_text(text)
    except Exception:
        # 如果 Markdown 标题切分失败，返回原始文本
        return [text]

    chunks = []
    for doc in docs:
        chunk_content = doc.page_content.strip()
        if not chunk_content:
            continue

        # 构建标题路径前缀
        header_parts = []
        for level in ["大章节", "小节", "小点", "段落"]:
            if level in doc.metadata and doc.metadata[level]:
                header_parts.append(doc.metadata[level])

        if header_parts:
            header_line = " > ".join(header_parts)
            chunk_content = f"{header_line}\n\n{chunk_content}"

        chunks.append(chunk_content)

    return chunks


def split_text_by_paragraphs(text: str, source_file: str = "") -> List[str]:
    """
    段落级切分：按标题 + 递归字符切分

    Args:
        text: 输入文本
        source_file: 来源文件名

    Returns:
        切分后的文本块列表
    """
    from langchain_text_splitters import RecursiveCharacterTextSplitter

    suffix = Path(source_file).suffix.lower() if source_file else ""

    # Markdown 文件优先按标题切分
    if suffix == ".md":
        header_chunks = split_text_by_headers(text)
        # 对过长的标题块再递归切分
        final_chunks = []
        for chunk in header_chunks:
            if len(chunk) <= settings.max_chunk_size:
                final_chunks.append(chunk)
            else:
                sub_splitter = RecursiveCharacterTextSplitter(
                    chunk_size=settings.max_chunk_size,
                    chunk_overlap=settings.chunk_overlap,
                    separators=["\n\n", "\n", "。", "！", "？", "；", "，", " ", ""],
                )
                final_chunks.extend(sub_splitter.split_text(chunk))
        return final_chunks

    # 其他格式：先尝试标题切分，再按字符长度切分
    try:
        from langchain_text_splitters import MarkdownHeaderTextSplitter
        header_splitter = MarkdownHeaderTextSplitter(
            headers_to_split_on=[
                ("#", "大章节"),
                ("##", "小节"),
                ("###", "小点"),
                ("####", "段落"),
            ],
            strip_headers=False,
        )
        header_chunks = header_splitter.split_text(text)
    except Exception:
        header_chunks = []
        # 如果标题切分失败，将整个文本作为一个 chunk
        from langchain_core.documents import Document
        header_chunks = [Document(page_content=text)]

    # 递归字符切分
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=settings.max_chunk_size,
        chunk_overlap=settings.chunk_overlap,
        separators=["\n\n", "\n", "。", "！", "？", "；", "，", " ", ""],
    )

    final_chunks = []
    for chunk in header_chunks:
        content = chunk.page_content.strip() if hasattr(chunk, 'page_content') else str(chunk).strip()
        if not content:
            continue

        if len(content) <= settings.max_chunk_size:
            final_chunks.append(content)
        else:
            sub_chunks = text_splitter.split_text(content)
            final_chunks.extend(sub_chunks)

    return final_chunks


def split_text(text: str, source_file: str = "") -> List[str]:
    """
    文本切片主入口

    流程：原始文本 → 清洗 → Markdown标题切分 → 递归字符切分 → 结果

    Args:
        text: 原始文本
        source_file: 来源文件名

    Returns:
        切分后的文本块列表
    """
    # 第一步：清洗
    cleaned = clean_text(text)

    if not cleaned:
        return []

    # 第二步：段落级切分
    chunks = split_text_by_paragraphs(cleaned, source_file)

    # 第三步：过滤太短的碎片
    final_chunks = [c for c in chunks if len(c.strip()) >= 10]

    return final_chunks


# ============================================================
#  第四部分：文件入库（整合入口）
# ============================================================

def ingest_file(file_path: Path) -> dict:
    """
    文件入库主流程

    流程：读取文件 → 文本提取 → 清洗 → 切片 → 向量嵌入 → 存入 ChromaDB

    Args:
        file_path: 文件路径

    Returns:
        {"file_name": ..., "chunks": ..., "chunk_texts": [...]}
    """
    from app.core.vector_store import get_vector_store

    file_name = file_path.name

    # 1. 提取文本
    raw_text = read_text_from_file(file_path)

    if not raw_text.strip():
        raise ValueError(f"文件 '{file_name}' 未提取到任何文本内容")

    # 2. 清洗 + 切片
    chunks = split_text(raw_text, file_name)

    if not chunks:
        raise ValueError(f"文件 '{file_name}' 切片后没有有效内容")

    # 3. 构建元数据
    metadatas = []
    for i, chunk in enumerate(chunks):
        metadatas.append({
            "source_file": file_name,
            "chunk_index": i,
            "chunk_size": len(chunk),
            "file_type": file_path.suffix.lower(),
        })

    # 4. 存入向量库（默认知识库）
    vector_store = get_vector_store()
    vector_store.add_texts(
        texts=chunks,
        metadatas=metadatas,
        source_file=file_name,
    )

    return {
        "file_name": file_name,
        "chunks": len(chunks),
        "chunk_texts": chunks,
    }


def ingest_file_to_collection(file_path: Path, collection_name: str) -> dict:
    """
    将文件入库到指定知识库

    Args:
        file_path: 文件路径
        collection_name: 目标知识库名称

    Returns:
        {"file_name": ..., "chunks": ..., "collection_name": ...}
    """
    from app.core.vector_store import get_vector_store

    file_name = file_path.name

    # 1. 提取文本
    raw_text = read_text_from_file(file_path)

    if not raw_text.strip():
        raise ValueError(f"文件 '{file_name}' 未提取到任何文本内容")

    # 2. 清洗 + 切片
    chunks = split_text(raw_text, file_name)

    if not chunks:
        raise ValueError(f"文件 '{file_name}' 切片后没有有效内容")

    # 3. 构建元数据
    metadatas = []
    for i, chunk in enumerate(chunks):
        metadatas.append({
            "source_file": file_name,
            "chunk_index": i,
            "chunk_size": len(chunk),
            "file_type": file_path.suffix.lower(),
        })

    # 4. 存入指定知识库
    vector_store = get_vector_store(collection_name)
    vector_store.add_texts(
        texts=chunks,
        metadatas=metadatas,
        source_file=file_name,
    )

    return {
        "file_name": file_name,
        "chunks": len(chunks),
        "collection_name": collection_name,
    }
